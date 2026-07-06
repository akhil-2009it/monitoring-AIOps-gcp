"""Generate a deep-dive Word doc that teaches AI / ML / MLOps / AIOps /
Data-Engineering concepts in order, stitches each topic to the next, and
maps every concept to a concrete file or resource in this repo.

Output: docs/Project_Learning_Guide.docx
Run:    python3 scripts/build_learning_doc.py
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "Project_Learning_Guide.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)


# ── helpers ─────────────────────────────────────────────────────────────
def shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def H(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = (
            RGBColor(0x0B, 0x4C, 0x8C) if level == 1 else
            RGBColor(0x1F, 0x6F, 0xB5) if level == 2 else
            RGBColor(0x44, 0x77, 0xAA)
        )
    return p


def P(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.bold = bold
    r.font.italic = italic
    return p


def CODE(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    return p


def BULLET(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(11)


def NUM(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for r in p.runs:
        r.font.size = Pt(11)


def CALLOUT(doc, title, body, color="DCE6F1"):
    """Single-row, single-cell shaded box with a title + body."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Light Grid Accent 1"
    cell = t.rows[0].cells[0]
    p1 = cell.paragraphs[0]
    r = p1.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x0B, 0x4C, 0x8C)
    p2 = cell.add_paragraph(body)
    for run in p2.runs:
        run.font.size = Pt(10)
    shade(cell, color)
    doc.add_paragraph()


def MAP_BOX(doc, concept, where, why):
    t = doc.add_table(rows=2, cols=3)
    t.style = "Light Grid Accent 1"
    head = t.rows[0].cells
    body = t.rows[1].cells
    head[0].text = "Concept"
    head[1].text = "Where in this repo"
    head[2].text = "Why it matters"
    for c in head:
        shade(c, "DCE6F1")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    body[0].text = concept
    body[1].text = where
    body[2].text = why
    for c in body:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
    doc.add_paragraph()


def STITCH(doc, prev, nxt):
    """Bridge paragraph that explicitly links chapter N to N+1."""
    box = doc.add_table(rows=1, cols=1)
    box.style = "Light Shading Accent 6"
    cell = box.rows[0].cells[0]
    shade(cell, "FFF4E5")
    p = cell.paragraphs[0]
    r = p.add_run("Stitch — what this hands off to the next chapter\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x99, 0x55, 0x00)
    r2 = p.add_run(f"You now have {prev}. The next chapter explains {nxt}.")
    r2.font.size = Pt(11)
    doc.add_paragraph()


def SECTION_BREAK(doc):
    p = doc.add_paragraph("─" * 70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


# ── content ─────────────────────────────────────────────────────────────
def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("monitoring-mlops-gcp\nLearning Guide — Deep Dive Edition")
    r.font.size = Pt(28); r.font.bold = True
    r.font.color.rgb = RGBColor(0x0B, 0x4C, 0x8C)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "AI · ML · MLOps · AIOps · Data Engineering · Cloud · Observability\n"
        "Each chapter taught from first principles and stitched into the next"
    )
    r.font.size = Pt(13); r.font.italic = True
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = intro.add_run(
        "\nThis guide is a single sequential thread. Read top to bottom. Every chapter "
        "ends with a 'Stitch' box that names exactly which problem the next chapter "
        "solves. Skip nothing — the order is the lesson."
    )
    r.font.size = Pt(11)

    doc.add_page_break()

    # ── Table of Contents ──────────────────────────────────────────────
    H(doc, "Table of contents", level=1)
    toc = [
        "1.  How software became impossible to watch by eye",
        "2.  Computers, the cloud, and what GCP is",
        "3.  Networking, identities, and why both matter before any data flows",
        "4.  Containers — the universal shipping unit",
        "5.  Kubernetes and GKE Autopilot — running many containers safely",
        "6.  Producers — applications that emit signals",
        "7.  The three pillars of observability — logs, metrics, traces",
        "8.  Pub/Sub — decoupling producers from consumers",
        "9.  Streaming ETL with Dataflow",
        "10. The data lake (GCS) and the warehouse (BigQuery)",
        "11. Schema and parsers — turning many formats into one",
        "12. Statistics 101 — what 'anomaly' actually means",
        "13. Streaming statistical detectors (Cloud Function Gen2)",
        "14. Machine Learning crash course",
        "15. Feature engineering — turning events into model inputs",
        "16. Training pipelines (Vertex Pipelines / KFP v2)",
        "17. Model registry and versioning",
        "18. Online endpoints (Vertex AI Endpoints)",
        "19. The Scoring API — wiring detectors into one HTTP surface",
        "20. The AIOps Console — the human-facing UI",
        "21. Drift, retraining, and the closed loop",
        "22. Security across the stack — defence in depth",
        "23. Cost engineering — knobs that decide your bill",
        "24. The full deploy walk-through (scripts/deploy_all.sh)",
        "25. Failure modes you will hit, and how to fix them",
        "26. Glossary",
        "27. A 4-week study plan",
    ]
    for line in toc:
        BULLET(doc, line)

    doc.add_page_break()

    # ── Chapter 1 ───────────────────────────────────────────────────────
    H(doc, "Chapter 1 — How software became impossible to watch by eye", level=1)

    H(doc, "1.1  The shape of a modern application", level=2)
    P(doc,
      "Twenty years ago a typical website was one binary on one server. One log file, one "
      "graph of CPU, one engineer on call. Today the same product is split into dozens of "
      "small services (microservices), each running multiple replicas, each emitting its own "
      "logs, metrics and traces. The volume of telemetry grows roughly with the square of "
      "team size — more services × more replicas × more requests."
    )

    H(doc, "1.2  Why humans cannot keep up", level=2)
    P(doc,
      "A single production cluster now produces millions of log lines per minute. Bugs hide "
      "between services. A single user request can hop through ten services and four databases "
      "before returning. Traditional dashboards (CPU > 80% → page someone) miss subtle issues: "
      "p99 latency rising, queue lag creeping up, a brand-new error message that has never "
      "appeared before but now appears once per second."
    )

    H(doc, "1.3  Three layers of help", level=2)
    NUM(doc, "Observability — collect logs, metrics, traces; let humans search.")
    NUM(doc, "Statistical alerting — define thresholds and rate-of-change rules; alert on breach.")
    NUM(doc, "AIOps — train models on historical signal; have them shout when something rare or unprecedented happens.")

    H(doc, "1.4  Why combine all three", level=2)
    P(doc,
      "Tier 1 catches the obvious; tier 2 catches the merely surprising; tier 3 catches the "
      "subtle. A serious platform layers all three. This repo is exactly that: managed GCP "
      "detectors (tier 1) + a Cloud Function with rules (tier 2) + four trained ML models "
      "(tier 3), funnelled into one alert UI."
    )

    CALLOUT(doc, "What is in this repo for chapter 1",
            "Read README.md, CLAUDE.md, and docs/DEPLOY.md. They describe the four-pillar "
            "architecture and the deploy steps. Everything below maps a piece of that "
            "architecture back to a concept.")

    STITCH(doc,
           "a clear picture of why a platform like this exists",
           "the underlying compute substrate — what 'the cloud' actually is, and which GCP "
           "primitives we will keep meeting in every later chapter")

    doc.add_page_break()

    # ── Chapter 2 — Cloud + GCP ─────────────────────────────────────────
    H(doc, "Chapter 2 — Computers, the cloud, and what GCP is", level=1)

    H(doc, "2.1  From physical servers to APIs", level=2)
    P(doc,
      "A cloud provider is a data centre exposed over HTTPS. Instead of buying and racking "
      "hardware, you call an API to get a VM, a database, a queue, or a managed service. "
      "Three big public providers exist: AWS (largest), Azure (Microsoft), and GCP (Google). "
      "All three offer the same primitives with different names and different defaults."
    )

    H(doc, "2.2  The four primitives every cloud provides", level=2)
    BULLET(doc, "Compute — VMs, containers, functions.")
    BULLET(doc, "Storage — object stores (file blobs), block stores (disks), databases (SQL/NoSQL).")
    BULLET(doc, "Networking — VPCs, load balancers, firewalls.")
    BULLET(doc, "Identity — service accounts, IAM roles, encryption keys.")

    H(doc, "2.3  Why GCP for this project", level=2)
    P(doc,
      "GCP shines for data-heavy ML workloads: BigQuery and Vertex AI are best-in-class. "
      "Pub/Sub is simpler than Kafka. Workload Identity is cleaner than AWS IRSA. The region "
      "we use is asia-south1 (Mumbai) — pick the region closest to your users."
    )

    H(doc, "2.4  The exact GCP services we will keep meeting", level=2)

    services = [
        ("Cloud Storage (GCS)", "S3 (AWS)", "Object store. Buckets full of files. Used for the data lake."),
        ("Pub/Sub", "Kinesis / Kafka", "Pub-sub message bus. Producers write, consumers read."),
        ("Dataflow", "Spark Streaming", "Managed Apache Beam — moves Pub/Sub → GCS at scale."),
        ("BigQuery", "Snowflake", "Serverless SQL warehouse. Query 100s of GB in seconds."),
        ("GKE Autopilot", "EKS / Kubernetes", "Managed Kubernetes. Runs our APIs as containers."),
        ("Vertex AI", "SageMaker", "Train + serve ML. Pipelines, Model Registry, Endpoints."),
        ("Cloud Functions Gen2", "Lambda", "Tiny serverless code triggered by events."),
        ("Cloud SQL", "RDS", "Managed Postgres / MySQL."),
        ("Memorystore", "ElastiCache", "Managed Redis."),
        ("Cloud Logging", "CloudWatch Logs", "Captures every container's stdout."),
        ("Cloud Trace", "X-Ray", "Distributed tracing."),
        ("Managed Prometheus (GMP)", "AMP", "Pulls /metrics endpoints, stores time-series."),
        ("Cloud Monitoring", "CloudWatch Metrics", "Dashboards + alert policies."),
        ("Cloud Armor", "WAF / Shield", "Web Application Firewall + DDoS rate limit."),
        ("Artifact Registry", "ECR", "Stores Docker images."),
        ("Secret Manager", "Secrets Manager", "Holds passwords/API keys."),
        ("IAM + Workload Identity", "IRSA on EKS", "Pods can act as a GCP identity safely."),
    ]
    t = doc.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
    head = t.rows[0].cells
    head[0].text, head[1].text, head[2].text = "GCP service", "Equivalent in AWS / OSS", "Job in this repo"
    for c in head:
        shade(c, "DCE6F1")
        for p in c.paragraphs:
            for r in p.runs: r.bold = True
    for s in services:
        row = t.add_row().cells
        for ci, val in enumerate(s):
            row[ci].text = val
            for p in row[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(10)

    H(doc, "2.5  Anatomy of a Terraform resource", level=2)
    P(doc,
      "Every GCP resource we create is declared as Terraform code. Terraform reads the current "
      "state, compares it to the code, and calls the GCP API to make reality match the code. "
      "Re-running Terraform with no code change is a no-op."
    )
    CODE(doc,
         'resource "google_storage_bucket" "lake" {\n'
         '  name                        = "monitoring-mlops-gcp-dev"\n'
         '  location                    = "asia-south1"\n'
         '  uniform_bucket_level_access = true\n'
         '  versioning { enabled = true }\n'
         '}')

    MAP_BOX(doc,
            "Terraform module pattern",
            "infra/main.tf composes infra/modules/{datalake,streaming,gke,...}/main.tf",
            "Splitting modules makes the codebase readable, testable, and reusable across "
            "environments (dev/stage/prod).")

    STITCH(doc,
           "a vocabulary of GCP services and a mental model of how Terraform creates them",
           "the two foundations every cloud workload sits on top of — networking and identity. "
           "Skip these two and nothing later will work")

    doc.add_page_break()

    # ── Chapter 3 — Networking + Identity ───────────────────────────────
    H(doc, "Chapter 3 — Networking, identities, and why both matter first", level=1)

    H(doc, "3.1  VPCs and the default network", level=2)
    P(doc,
      "A VPC (Virtual Private Cloud) is a private network in GCP. Resources you put inside "
      "share a private IP space and can talk without crossing the public internet. Every GCP "
      "project starts with a 'default' network — fine for learning, replaced by custom VPCs "
      "in real prod. This project uses default."
    )

    H(doc, "3.2  Private Service Access (PSA)", level=2)
    P(doc,
      "Cloud SQL and Memorystore live in Google-managed networks. To talk to them privately "
      "from your VPC, you reserve an IP range and peer it via the Service Networking API. "
      "That is what google_compute_global_address.psa + google_service_networking_connection.psa "
      "do in modules/database/main.tf."
    )

    H(doc, "3.3  Static IPs + load balancers", level=2)
    P(doc,
      "External users hit a global static IP. The IP is forwarded to a Google global HTTPS "
      "load balancer, which routes to a backend (a GKE Service in our case). Static IPs are "
      "named so DNS records stay stable across redeploys: aiops-scoring-ip, aiops-ui-ip, "
      "demo-app-ip."
    )

    H(doc, "3.4  IAM — three identity types", level=2)
    BULLET(doc, "User accounts — humans (owner@example.com).")
    BULLET(doc, "Service accounts — robots (aiops-runner@PROJECT.iam.gserviceaccount.com).")
    BULLET(doc, "Workload Identity — a GKE pod acts as a service account without static keys.")

    H(doc, "3.5  Roles vs permissions", level=2)
    P(doc,
      "A permission is one tiny capability (e.g. storage.objects.create). A role is a bag of "
      "permissions (e.g. roles/storage.objectAdmin). You bind a role to a principal on a "
      "scope (project, bucket, secret). Grant the smallest scope that works — 'least privilege'."
    )

    H(doc, "3.6  Workload Identity end-to-end", level=2)
    NUM(doc, "Cluster has a Workload Identity Pool: PROJECT.svc.id.goog.")
    NUM(doc, "Each Kubernetes Service Account (KSA) maps to a GCP Service Account (GSA) via an annotation.")
    NUM(doc, "GSA grants iam.workloadIdentityUser to KSA → pod gets short-lived OIDC token.")
    NUM(doc, "Token is exchanged for a real GCP access token at runtime by the metadata server inside the pod.")

    MAP_BOX(doc,
            "WLI in this repo",
            "infra/modules/gke/main.tf — google_service_account_iam_member.wli_scoring / wli_otel / wli_ui. "
            "demo-app/infra/main.tf — google_service_account_iam_member.demo_api_wi / demo_worker_wi.",
            "Every pod that talks to a GCP API uses WLI. Zero static credentials in the cluster.")

    H(doc, "3.7  Cloud Armor — perimeter security", level=2)
    P(doc,
      "Cloud Armor is a WAF + DDoS rate-limiter that sits in front of the global load balancer. "
      "We attach it via a BackendConfig CRD on the Service. Two rule families: "
      "rate_limit (600 req/min/IP) and pre-built WAF rulesets (sqli-v33-stable, xss-v33-stable)."
    )

    STITCH(doc,
           "a network the workloads can sit in and an identity model for them",
           "the lowest-level unit a workload is shipped as — the container — and the format "
           "that defines one (Dockerfile)")

    doc.add_page_break()

    # ── Chapter 4 — Containers ──────────────────────────────────────────
    H(doc, "Chapter 4 — Containers — the universal shipping unit", level=1)

    H(doc, "4.1  Why containers", level=2)
    P(doc,
      "Before containers, deploying meant copying code to a server, hoping its libraries were "
      "the same version as on your laptop, and patching when they weren't. A container is a "
      "tar file that bundles the code, the libraries, and even the OS user-space — everything "
      "but the kernel. Docker is the most common container format and runtime."
    )

    H(doc, "4.2  Anatomy of a Dockerfile", level=2)
    CODE(doc,
         "FROM python:3.11-slim\n"
         "ENV PYTHONUNBUFFERED=1\n"
         "WORKDIR /app\n"
         "COPY requirements.txt /tmp/\n"
         "RUN pip install -r /tmp/requirements.txt\n"
         "COPY app /app/app\n"
         'CMD ["uvicorn", "app.main:app", "--host", "0.0.0.0", "--port", "8000"]')

    BULLET(doc, "FROM — base image. Layers stack on top of layers.")
    BULLET(doc, "RUN — execute a build-time command. Cached unless inputs change.")
    BULLET(doc, "COPY — pull files from build context into the image.")
    BULLET(doc, "CMD — what to run when a container starts. One process, foreground.")

    H(doc, "4.3  Image registries", level=2)
    P(doc,
      "Built images are pushed to a registry — a content-addressable store of layers. GCP's "
      "registry is Artifact Registry. Each image has a tag (often the git SHA + 'latest'). "
      "Pulling an image by tag downloads only the layers you don't already have."
    )

    H(doc, "4.4  Build best practices we follow", level=2)
    BULLET(doc, "Multi-stage builds for compiled languages — small final image.")
    BULLET(doc, "Pin the base image tag (no 'latest' in the FROM line).")
    BULLET(doc, "Run as non-root (the api Dockerfile creates 'appuser').")
    BULLET(doc, "Add a HEALTHCHECK so orchestrators can self-heal.")
    BULLET(doc, "Keep layers small — install deps before copying code so a code change doesn't bust the deps cache.")

    MAP_BOX(doc,
            "Dockerfiles in this repo",
            "api/scoring/Dockerfile — Scoring API\n"
            "api/ui/Dockerfile — AIOps Console (nginx + static)\n"
            "demo-app/api/Dockerfile, demo-app/worker/Dockerfile, demo-app/web/Dockerfile\n"
            "ml/pipelines/rcf_metrics/Dockerfile — trainer image",
            "Each is built + pushed by Stage 2 of scripts/deploy_all.sh.")

    STITCH(doc,
           "an image that holds one process and its dependencies",
           "the system that runs hundreds of those images, schedules them across machines, "
           "restarts crashed ones, and exposes them on the network — Kubernetes / GKE")

    doc.add_page_break()

    # ── Chapter 5 — Kubernetes / GKE ───────────────────────────────────
    H(doc, "Chapter 5 — Kubernetes and GKE Autopilot — running many containers safely", level=1)

    H(doc, "5.1  Why an orchestrator", level=2)
    P(doc,
      "If you have ten machines and twenty container types, you do not want to ssh and 'docker "
      "run' by hand. Kubernetes (k8s) runs containers across a fleet, restarts crashed ones, "
      "scales them, exposes them as a virtual IP, and spreads them across failure zones."
    )

    H(doc, "5.2  Core k8s objects you must know", level=2)
    BULLET(doc, "Pod — one or more containers sharing a network namespace. Smallest unit.")
    BULLET(doc, "Deployment — declares 'I want N replicas of this pod template'.")
    BULLET(doc, "Service — stable virtual IP + DNS for a set of pods.")
    BULLET(doc, "ConfigMap / Secret — env vars and credentials, mounted into pods.")
    BULLET(doc, "Ingress — HTTP routing rules from the outside world to a Service.")
    BULLET(doc, "ServiceAccount — the pod's identity (used by Workload Identity in GKE).")
    BULLET(doc, "HPA — HorizontalPodAutoscaler, scales replicas on CPU / custom metrics.")
    BULLET(doc, "PodDisruptionBudget — prevents the cluster from evicting too many pods at once.")
    BULLET(doc, "NetworkPolicy — pod-level firewall rules.")
    BULLET(doc, "CronJob — runs a Job on a schedule.")

    H(doc, "5.3  GKE Autopilot specifically", level=2)
    P(doc,
      "Autopilot is a GKE mode where Google manages the underlying nodes. You only declare "
      "pods. Cluster scales in 1-2 minutes; you pay per pod-second. Constraints: certain "
      "privileged options are blocked, but everything we use here works."
    )

    H(doc, "5.4  Helm — the package manager", level=2)
    P(doc,
      "A Helm 'chart' is a templated bundle of YAML. values.yaml fills the template; you "
      "install with 'helm upgrade --install'. Helm is to Kubernetes what apt is to Debian."
    )

    H(doc, "5.5  Charts in this repo (read each templates/ folder)", level=2)
    BULLET(doc, "helm/charts/anomaly-scoring-api — the FastAPI scoring service.")
    BULLET(doc, "helm/charts/aiops-ui — the static web console.")
    BULLET(doc, "helm/charts/otel-collector — receives OTLP traces.")
    BULLET(doc, "demo-app/helm/demo-service — generic chart, parametrised per service via values-{api,worker,web}.yaml.")

    H(doc, "5.6  How traffic enters the cluster", level=2)
    CODE(doc,
         "Internet ─► Global LB (with static IP + Cloud Armor) ─► Ingress\n"
         "                                                            │\n"
         "                                                            ▼\n"
         "                                                       Service ─► Pod(s)")

    MAP_BOX(doc,
            "Ingress definitions",
            "helm/charts/aiops-ui/templates/ingress.yaml + service.yaml — UI Ingress, ManagedCertificate, BackendConfig\n"
            "helm/charts/anomaly-scoring-api/templates/ingress.yaml — API Ingress\n"
            "demo-app/helm/demo-service/templates/ingress.yaml — demo Ingress",
            "These three charts are the only ones with a public face. Everything else is ClusterIP.")

    STITCH(doc,
           "a cluster that runs containers and exposes them on the network",
           "the workloads we put inside it — applications that emit signals (logs / metrics / "
           "traces). Without those signals there is nothing for the AIOps half of the project to do")

    doc.add_page_break()

    # ── Chapter 6 — Producers ───────────────────────────────────────────
    H(doc, "Chapter 6 — Producers — applications that emit signals", level=1)

    H(doc, "6.1  Why we built a fake e-commerce stack", level=2)
    P(doc,
      "An anomaly-detection platform needs raw signal. demo-app/ exists so you can run an "
      "end-to-end demo without pointing the platform at real production. It has three roles "
      "and a traffic generator."
    )

    H(doc, "6.2  The three demo services", level=2)
    BULLET(doc, "demo-app/web — NGINX reverse proxy + static index.html. Combined access log.")
    BULLET(doc, "demo-app/api — FastAPI backend. JSON logs, /metrics, OTEL traces, talks to MySQL.")
    BULLET(doc, "demo-app/worker — Python loop. Polls Redis queue, processes synthetic jobs.")

    H(doc, "6.3  Realism knobs", level=2)
    P(doc,
      "demo-app/api/app/main.py + routes.py read environment variables that inject controlled "
      "chaos so detectors have something to fire on:"
    )
    BULLET(doc, "DEMO_ERROR_RATE — fraction of requests that deliberately 500.")
    BULLET(doc, "DEMO_SLOW_RATE — fraction of requests that get an extra sleep.")
    BULLET(doc, "DEMO_SLOW_LATENCY — how much extra sleep.")

    H(doc, "6.4  Traffic generator (Locust)", level=2)
    P(doc,
      "demo-app/traffic-gen/locustfile.py defines normal-mode and attack-mode user behaviours "
      "(brute-force /login, DDoS bursts, SQLi probes). A Kubernetes CronJob "
      "(demo-app/traffic-gen/k8s-cronjob.yaml) re-spawns it every 30 minutes."
    )

    MAP_BOX(doc,
            "x-cocoon-id propagation",
            "demo-app/{api,worker,web} — middleware reads the x-cocoon-id header and echoes it",
            "Cocoons are ephemeral isolated test environments (from the wider Skillz Cocoon "
            "platform). Out of scope here, but the pattern is real-world useful: tag every "
            "log with a request-scoped trace key.")

    STITCH(doc,
           "applications that produce structured signal continuously",
           "the three formats that signal travels in — logs, metrics, traces — and how each "
           "is captured and stored")

    doc.add_page_break()

    # ── Chapter 7 — Observability ───────────────────────────────────────
    H(doc, "Chapter 7 — The three pillars of observability", level=1)

    H(doc, "7.1  Logs", level=2)
    P(doc,
      "Discrete time-stamped events. JSON lines on stdout. GKE captures them automatically "
      "into Cloud Logging. Searchable, filterable, exportable to BigQuery / GCS / Pub/Sub via "
      "log sinks."
    )
    CODE(doc,
         '{"ts":"2026-06-20T09:30:14Z","level":"INFO","msg":"GET /api/products/5 -> 200 (12.3ms)",'
         '"service":"demo-api","request_id":"u-7f","trace_id":"a91…"}')

    H(doc, "7.2  Metrics", level=2)
    P(doc,
      "Numerical time-series. Three shapes:"
    )
    BULLET(doc, "Counter — only goes up. e.g. demo_api_requests_total.")
    BULLET(doc, "Gauge — goes up and down. e.g. demo_api_inflight_requests.")
    BULLET(doc, "Histogram — bucketed distribution. e.g. demo_api_request_latency_seconds.")
    P(doc,
      "Apps expose /metrics in Prometheus text format. Google Managed Prometheus (GMP) scrapes "
      "them via PodMonitoring CRDs. Histograms are the secret weapon — you can compute p50, "
      "p95, p99 on the fly server-side."
    )

    H(doc, "7.3  Traces", level=2)
    P(doc,
      "A trace is a tree of spans. Each span = one operation in one service ('SQL SELECT', "
      "'POST /score', 'redis.lpop'). Spans share a trace_id so you can replay a request that "
      "hopped api → worker → MySQL → Redis as one timeline. OpenTelemetry is the vendor-neutral "
      "SDK. Apps emit OTLP; the OTEL collector exports to Cloud Trace."
    )

    MAP_BOX(doc,
            "Where each pillar is wired",
            "Logs — JsonFormatter in demo-app/api/app/observability.py + GKE auto-capture\n"
            "Metrics — prometheus_client Counter/Histogram/Gauge in same file + PodMonitoring CR\n"
            "Traces — OpenTelemetry SDK exporters → otel-collector → Cloud Trace",
            "All three share the request_id and trace_id so logs and traces correlate.")

    H(doc, "7.4  Why all three matter", level=2)
    BULLET(doc, "Metrics — fastest to graph and alert on. Cheapest to store.")
    BULLET(doc, "Logs — best for forensics ('what exactly did the broken request look like?').")
    BULLET(doc, "Traces — best for debugging slowness across many services.")

    STITCH(doc,
           "three streams of telemetry being emitted by every workload",
           "the bus that moves them at scale and decouples producers from consumers — Pub/Sub")

    doc.add_page_break()

    # ── Chapter 8 — Pub/Sub ─────────────────────────────────────────────
    H(doc, "Chapter 8 — Pub/Sub — decoupling producers from consumers", level=1)

    H(doc, "8.1  The publish / subscribe pattern", level=2)
    P(doc,
      "Producers send messages to a topic. Consumers create a subscription on the topic and "
      "pull messages. Producers do not know who consumes. Consumers do not know who produced. "
      "Adding a new consumer is one new subscription — no producer change."
    )

    H(doc, "8.2  Three topics in this repo", level=2)
    t = doc.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
    head = t.rows[0].cells
    for ci, h in enumerate(["Topic", "Producers", "Consumers"]):
        head[ci].text = h; shade(head[ci], "DCE6F1")
        for p in head[ci].paragraphs:
            for r in p.runs: r.bold = True
    rows = [
        ("monitoring-mlops-{env}-events",
         "Cloud Logging sinks (apps, NGINX, LB, Cloud SQL)",
         "Streaming Cloud Function (rules) + Dataflow GCS writer"),
        ("monitoring-mlops-{env}-anomalies",
         "Streaming Function + Scoring API",
         "Anomaly fan-out subscription (UI / Slack / ticketing)"),
        ("monitoring-mlops-{env}-retrain",
         "Cloud Scheduler",
         "Pipeline-trigger Cloud Function"),
    ]
    for row in rows:
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph()

    H(doc, "8.3  Subscriptions and ack-deadline", level=2)
    P(doc,
      "Each consumer needs its own subscription. When you pull a message the broker starts a "
      "timer (ack_deadline_seconds) — you must ack before it expires or it is redelivered. "
      "Long jobs need long ack deadlines. Streaming function = 60s; Dataflow writer = 600s."
    )

    H(doc, "8.4  Cloud Logging sinks — the entry point", level=2)
    P(doc,
      "Cloud Logging captures every container's stdout. A 'sink' is a saved filter that "
      "routes matching log lines to GCS, BigQuery or Pub/Sub. We use Pub/Sub as the routing "
      "fabric so all downstream paths read from one bus."
    )
    MAP_BOX(doc,
            "Sink config",
            "demo-app/infra/main.tf — google_logging_project_sink.mysql_slow_sink",
            "Mysql slow query log → events Pub/Sub topic. Add more sinks for app/LB logs the same way.")

    STITCH(doc,
           "a fan-in bus that holds every signal",
           "the streaming ETL job that drains the bus into a long-term store — Dataflow")

    doc.add_page_break()

    # ── Chapter 9 — Dataflow ────────────────────────────────────────────
    H(doc, "Chapter 9 — Streaming ETL with Dataflow", level=1)

    H(doc, "9.1  ETL — what those three letters mean", level=2)
    BULLET(doc, "Extract — read from source (Pub/Sub).")
    BULLET(doc, "Transform — reshape (parse JSON, derive fields, tag with partition).")
    BULLET(doc, "Load — write to sink (GCS files in our case).")

    H(doc, "9.2  Apache Beam + Dataflow", level=2)
    P(doc,
      "Apache Beam is a unified programming model for batch and streaming data pipelines. "
      "Dataflow is Google's managed runner for Beam — autoscales workers, handles late "
      "data, exactly-once semantics. We do not write Beam code; we use Google's pre-built "
      "PubSub-to-GCS-Text flex template."
    )

    H(doc, "9.3  Windowing", level=2)
    P(doc,
      "Streaming jobs accumulate messages into windows before writing. We use a 5-minute "
      "fixed window: every 5 minutes the job flushes a file to GCS. Smaller windows = more "
      "files = more overhead. Bigger windows = more lag before data appears in the lake."
    )

    H(doc, "9.4  Hive partitioning", level=2)
    P(doc,
      "We write to gs://bucket/{env}/raw/source=app/...,  source=lb/..., source=cloudsql/...  "
      "BigQuery, Spark, and most query engines understand this directory layout — they can "
      "skip files when the WHERE clause filters by source."
    )

    MAP_BOX(doc,
            "Dataflow job in this repo",
            "infra/modules/streaming/main.tf — google_dataflow_flex_template_job.events_to_gcs",
            "Reads events_to_dataflow subscription, writes 5-minute files into the lake bucket "
            "under {env}/raw/.")

    STITCH(doc,
           "a process that durably stores every signal",
           "the storage layout itself — the data lake (raw) plus the warehouse view (BigQuery)")

    doc.add_page_break()

    # ── Chapter 10 — Lake + Warehouse ───────────────────────────────────
    H(doc, "Chapter 10 — The data lake (GCS) and the warehouse (BigQuery)", level=1)

    H(doc, "10.1  Lake vs warehouse", level=2)
    BULLET(doc, "Lake — raw, schema-on-read, cheap, infinite. Files in GCS.")
    BULLET(doc, "Warehouse — schema-on-write, query-optimised, expensive per query. BigQuery.")
    P(doc, "Modern systems use both. The lake is the source of truth; the warehouse is the analyst's playground.")

    H(doc, "10.2  Bucket layout we use", level=2)
    CODE(doc,
         "gs://monitoring-mlops-gcp-dev/\n"
         "  dev/raw/source=app/<datepartitions>/events-*.json\n"
         "  dev/raw/source=lb/...\n"
         "  dev/raw/source=cloudsql/...\n"
         "  dev/features/security/<datepartitions>/*.parquet\n"
         "  dev/dataflow/tmp/...\n"
         "  dev/models/<detector>/<run-id>/...")

    H(doc, "10.3  External tables", level=2)
    P(doc,
      "An external table is a BigQuery view over GCS files. No data is copied. Schema is "
      "auto-detected. With hive partitioning, queries with WHERE source='app' only scan files "
      "under source=app/."
    )

    H(doc, "10.4  Native tables", level=2)
    P(doc,
      "A native table physically stores data in BigQuery's own storage — fast aggregation, "
      "DAY partitioning, clustering. We use one for anomalies because we want sub-second "
      "queries from the dashboard."
    )

    MAP_BOX(doc,
            "Datalake module",
            "infra/modules/datalake/main.tf — bucket + bigquery_dataset 'monitoring' + 3 tables",
            "Tables: raw_events (external, hive-partitioned by source), features_security "
            "(external Parquet), anomalies (native, DAY-partitioned).")

    STITCH(doc,
           "raw data on disk and queryable in SQL",
           "the schema discipline that lets many sources end up in one table — parsers and "
           "the CommonEvent format")

    doc.add_page_break()

    # ── Chapter 11 — Schema + parsers ──────────────────────────────────
    H(doc, "Chapter 11 — Schema and parsers — turning many formats into one", level=1)

    H(doc, "11.1  Why a canonical schema", level=2)
    P(doc,
      "Each source emits a different shape: NGINX combined log, MySQL slow query, OTEL JSON, "
      "Cloud Armor. If detectors had to know every shape, adding a source would touch every "
      "detector. Solution: one CommonEvent schema; per-source parsers that adapt. Detectors "
      "read CommonEvent only."
    )

    H(doc, "11.2  CommonEvent shape", level=2)
    CODE(doc,
         "{\n"
         "  ts:        iso8601,\n"
         "  ingest_ts: iso8601,\n"
         "  source:    'cdn|lb|cloud_armor|app|gke|nginx|cloudsql|...',\n"
         "  host:      str,\n"
         "  severity:  'DEBUG|INFO|WARN|ERROR|CRITICAL',\n"
         "  status:    int|None,\n"
         "  latency_ms: float|None,\n"
         "  bytes:     int|None,\n"
         "  src_ip:    str|None,\n"
         "  user:      str|None,\n"
         "  path:      str|None,\n"
         "  user_agent: str|None,\n"
         "  message:   str,\n"
         "  attrs:     dict\n"
         "}")

    H(doc, "11.3  PII handling", level=2)
    P(doc,
      "CLAUDE.md rule 1: hash usernames + IPs (HMAC) inside the parser. Detectors see hashes "
      "only. Lets you analyse rate per user without ever storing or logging user identifiers."
    )

    MAP_BOX(doc,
            "Parsers",
            "ml/parsers/__init__.py — CommonEvent dataclass + parse() dispatcher per source",
            "Add a new source = write one parse_<source>() function and add to the dispatcher.")

    STITCH(doc,
           "data normalised to one shape",
           "the math that decides which CommonEvent rows are 'anomalous'")

    doc.add_page_break()

    # ── Chapter 12 — Stats 101 ─────────────────────────────────────────
    H(doc, "Chapter 12 — Statistics 101 — what 'anomaly' actually means", level=1)

    H(doc, "12.1  Mean and standard deviation", level=2)
    P(doc,
      "Mean (μ) — the centre. Standard deviation (σ) — typical distance from centre. "
      "If a metric's recent mean is 100 req/s with σ=15, then 250 req/s is (250-100)/15 = 10 "
      "standard deviations away — extremely unlikely under normal behaviour."
    )

    H(doc, "12.2  Z-score", level=2)
    P(doc, "z = (x - μ) / σ. |z| > 3 is a common 'anomaly' threshold. Cheap, robust on stationary data.")

    H(doc, "12.3  EWMA — exponentially weighted moving average", level=2)
    P(doc,
      "Plain mean treats yesterday and last week the same. EWMA weights recent points more: "
      "EWMA_t = α·x_t + (1-α)·EWMA_{t-1}. Adapts to drifts."
    )

    H(doc, "12.4  Rate-of-change", level=2)
    P(doc, "If a counter doubles in 30 seconds, that itself is news. Rate-of-change rules "
            "catch sudden spikes that z-score on raw values misses if the baseline window "
            "is too long.")

    H(doc, "12.5  Distinct counters", level=2)
    P(doc,
      "How many unique src_ips in the last minute? A jump from 50 to 5000 distinct IPs is the "
      "signature of a DDoS even before total request rate explodes. Implemented with HyperLogLog "
      "or a bounded set."
    )

    H(doc, "12.6  Why these still aren't enough", level=2)
    BULLET(doc, "Seasonality — traffic is not stationary; lunch peak ≠ midnight.")
    BULLET(doc, "Multi-variate — error rate alone is fine; error rate AND latency rising AND "
                "queue depth growing means something specific is breaking.")
    BULLET(doc, "Cold start — z-score on the first 30 minutes is meaningless; ML detectors trained "
                "on weeks of data win the long tail.")

    STITCH(doc,
           "the basic statistical tools of anomaly detection",
           "how those tools are wired into a live, sub-second alerting pipeline — the streaming "
           "Cloud Function")

    doc.add_page_break()

    # ── Chapter 13 — Streaming detector ────────────────────────────────
    H(doc, "Chapter 13 — Streaming statistical detectors (Cloud Function Gen2)", level=1)

    H(doc, "13.1  Cloud Functions Gen2", level=2)
    P(doc,
      "A managed runtime that executes a small piece of code in response to an event "
      "(here, a Pub/Sub message). Auto-scales from 0 to thousands of concurrent instances. "
      "Uses Cloud Run under the hood — 1 GB RAM, 120 s timeout in our config."
    )

    H(doc, "13.2  Entry point", level=2)
    CODE(doc,
         "@functions_framework.cloud_event\n"
         "def handler(event):\n"
         "    msg = base64.b64decode(event.data['message']['data']).decode()\n"
         "    record = json.loads(msg)\n"
         "    for rule in _rules():\n"
         "        if _check(rule, record['source'], record['host'], record['value']):\n"
         "            _emit(...)")

    H(doc, "13.3  Rules", level=2)
    P(doc,
      "ml/streaming/rules.yaml is plain YAML — easier to tune than code. Each rule names a "
      "metric, a method (z-score / ewma / rate / distinct), thresholds, and severity. "
      "The function loads the YAML at cold start; on a hot instance it stays resident."
    )

    H(doc, "13.4  Cold-state caveat", level=2)
    P(doc,
      "A new function instance has no rolling history. The first ~30 minutes of data after a "
      "scale-out may produce noisy z-scores. Acceptable trade-off because the function is "
      "dirt cheap and ML detectors cover that gap."
    )

    H(doc, "13.5  Where anomalies go", level=2)
    BULLET(doc, "Pub/Sub topic — anomalies fan-out subscription.")
    BULLET(doc, "Cloud Logging — severity from rule, indexed by detector + source.")
    BULLET(doc, "Cloud Monitoring custom metric — aiops/streaming/anomalies counter.")

    MAP_BOX(doc,
            "Streaming detector code",
            "ml/streaming/detector.py + rules.yaml",
            "Deployed by Stage 4 of scripts/deploy_all.sh.")

    STITCH(doc,
           "a tier-2 detector that fires within seconds without any training",
           "the foundation needed for tier-3 — what machine learning is, in just enough depth "
           "to follow the rest of the project")

    doc.add_page_break()

    # ── Chapter 14 — ML crash course ───────────────────────────────────
    H(doc, "Chapter 14 — Machine Learning crash course", level=1)

    H(doc, "14.1  What is a model", level=2)
    P(doc,
      "A function f(x)→y whose parameters are fit from data. Training = pick parameters that "
      "minimise loss on a labelled (or unlabelled) dataset. Inference = call f on new x."
    )

    H(doc, "14.2  Three paradigms", level=2)
    BULLET(doc, "Supervised — y is given. Predict y from x. (spam classifier)")
    BULLET(doc, "Unsupervised — no y. Find structure (cluster, density). (anomaly detection)")
    BULLET(doc, "Self-supervised — make y from x (mask a word, predict it). (BERT, GPT pre-training)")

    H(doc, "14.3  Train / validate / test", level=2)
    P(doc,
      "Split data three ways. Train on 'train'. Tune hyperparameters on 'validation'. Estimate "
      "real-world performance on 'test' (touched once at the end). This avoids looking good "
      "on data the model has memorised."
    )

    H(doc, "14.4  Metrics", level=2)
    BULLET(doc, "Accuracy — correct / total. Bad on imbalanced classes.")
    BULLET(doc, "Precision — of predicted positives, how many real. (Don't page on noise.)")
    BULLET(doc, "Recall — of real positives, how many caught. (Don't miss real outages.)")
    BULLET(doc, "F1 — harmonic mean of precision and recall.")
    BULLET(doc, "AUC — area under ROC curve. Threshold-independent ranking quality.")
    BULLET(doc, "P@1% — precision at the top 1% of scores. The detector-eval favourite.")

    H(doc, "14.5  Our four detectors", level=2)
    t = doc.add_table(rows=1, cols=4); t.style = "Light Grid Accent 1"
    head = t.rows[0].cells
    for ci, h in enumerate(["Detector", "Algorithm", "Strength", "Quality gate"]):
        head[ci].text = h; shade(head[ci], "DCE6F1")
        for p in head[ci].paragraphs:
            for r in p.runs: r.bold = True
    rows = [
        ("rcf_metrics", "Random Cut Forest", "Time-series numerical metrics", "F1 ≥ 0.70"),
        ("iforest_logs", "Isolation Forest", "Tabular features over log windows", "P@1% ≥ 0.80"),
        ("lstm_ae_traces", "LSTM Auto-Encoder", "Trace-sequence reconstruction error", "AUC > 0.80"),
        ("log_embedding_anomaly", "Log-BERT (transformer)", "Free-text app log lines", "P@1% ≥ 0.75"),
    ]
    for r in rows:
        cells = t.add_row().cells
        for ci, val in enumerate(r):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for run in p.runs: run.font.size = Pt(10)
    doc.add_paragraph()

    STITCH(doc,
           "a vocabulary for ML and the four specific algorithms in use",
           "the first concrete step in any ML pipeline — turning raw events into model inputs "
           "(feature engineering)")

    doc.add_page_break()

    # ── Chapter 15 — Feature engineering ───────────────────────────────
    H(doc, "Chapter 15 — Feature engineering — turning events into model inputs", level=1)

    H(doc, "15.1  Why features matter more than algorithm choice", level=2)
    P(doc,
      "Two competent ML practitioners with the same algorithm will produce wildly different "
      "results based on features. Garbage in, garbage out — this is the most-cited rule in "
      "applied ML for a reason."
    )

    H(doc, "15.2  Sliding windows", level=2)
    P(doc,
      "Raw events are point-in-time. Most useful features are over a window: requests in last "
      "5 min, distinct IPs in last 5 min, p95 latency in last 5 min. We use 5-minute fixed "
      "windows here. Rolling windows (overlapping) are smoother but ~5x more expensive."
    )

    H(doc, "15.3  Security features we compute", level=2)
    BULLET(doc, "req_count_5m — total requests by source / host.")
    BULLET(doc, "err_count_5m / err_rate_5m — 4xx + 5xx counts and ratio.")
    BULLET(doc, "p95_latency_5m — 95th percentile latency.")
    BULLET(doc, "distinct_src_ips_5m — unique source IPs.")
    BULLET(doc, "auth_failure_count_5m — count of 401 / 403 / login=failed.")
    BULLET(doc, "ua_entropy_5m — Shannon entropy of user-agent string distribution.")
    BULLET(doc, "path_diversity_5m — count of distinct request paths.")

    MAP_BOX(doc,
            "Feature builder",
            "ml/feature_engineering/security_features.py",
            "Reads raw events from gs://.../{env}/raw/, computes the 5-min features, writes "
            "Parquet to gs://.../{env}/features/security/. The four trainers point at this prefix.")

    STITCH(doc,
           "model-ready feature tables on disk",
           "the orchestration layer that turns those features into a trained, registered model "
           "— Vertex Pipelines")

    doc.add_page_break()

    # ── Chapter 16 — Pipelines ─────────────────────────────────────────
    H(doc, "Chapter 16 — Training pipelines (Vertex Pipelines / KFP v2)", level=1)

    H(doc, "16.1  Why a pipeline framework", level=2)
    P(doc,
      "A real training run has many steps: load → split → train → evaluate → register → "
      "(maybe) deploy. Doing them in a Jupyter notebook is fine for prototypes; in production "
      "you need versioned, parameterised, reproducible DAGs that record their inputs, outputs, "
      "and metrics. Vertex Pipelines (KFP v2 SDK) is Google's offering."
    )

    H(doc, "16.2  Components", level=2)
    P(doc,
      "A component is one Python function decorated with @kfp.dsl.component. Inputs and "
      "outputs are typed (Datasets, Models, Metrics). The KFP compiler builds a YAML and "
      "uploads it to Vertex; Vertex schedules each component as a containerised step."
    )

    H(doc, "16.3  Lineage", level=2)
    P(doc,
      "Each step's outputs are stored. Each model version is linked to the run that produced "
      "it, the data slices used, the metrics observed. When a detector behaves badly in prod "
      "you can answer 'what training run made this model?' in two clicks."
    )

    MAP_BOX(doc,
            "Our four pipelines",
            "ml/pipelines/{rcf_metrics,iforest_logs,lstm_ae_traces,log_embedding_anomaly}/pipeline.py",
            "Each defines load → train → eval → register. Trainer images live next to them.")

    H(doc, "16.4  Compute infra per pipeline", level=2)
    BULLET(doc, "RCF / IForest — n1-standard-4 Spot, ≤ 30 min, ~$0.05/run.")
    BULLET(doc, "LSTM-AE / LogBERT — g2-standard-8 + 1× L4 GPU Spot, 1-3 h, ~$0.30/run.")
    BULLET(doc, "Spot = preemptible. 3-4× cheaper. CLAUDE.md rule 2.")

    STITCH(doc,
           "trained model artefacts in cloud storage with metadata",
           "the catalogue that makes those artefacts discoverable, comparable and rollback-able "
           "— the model registry")

    doc.add_page_break()

    # ── Chapter 17 — Registry ──────────────────────────────────────────
    H(doc, "Chapter 17 — Model registry and versioning", level=1)

    H(doc, "17.1  What a registry holds", level=2)
    BULLET(doc, "Display name (e.g. rcf-metrics-detector).")
    BULLET(doc, "Version (auto-incremented per upload).")
    BULLET(doc, "Container image used to serve.")
    BULLET(doc, "Eval metrics (linked from the run).")
    BULLET(doc, "State labels (champion / challenger / archive).")

    H(doc, "17.2  Promotion flow", level=2)
    NUM(doc, "Pipeline runs nightly; if metrics pass the quality gate it uploads a new version.")
    NUM(doc, "Operator (or auto-policy) sets new version as 'default' on the endpoint with traffic split (10% then 100%).")
    NUM(doc, "Old version stays for instant rollback.")

    H(doc, "17.3  Why versioning matters operationally", level=2)
    P(doc,
      "The day a detector pages on-call uselessly, you do not want to retrain in a hurry. "
      "You want to roll back to v(N-1) in five seconds. That is a one-line gcloud call, "
      "thanks to the registry."
    )

    MAP_BOX(doc,
            "Registry in this repo",
            "Vertex AI Model Registry (managed). Pipelines call aiplatform.Model.upload(...).",
            "View at console.cloud.google.com/vertex-ai/models")

    STITCH(doc,
           "a versioned catalogue of trained detectors",
           "the live HTTPS endpoints that actually run inference at request time — Vertex AI "
           "Endpoints")

    doc.add_page_break()

    # ── Chapter 18 — Endpoints ─────────────────────────────────────────
    H(doc, "Chapter 18 — Online endpoints (Vertex AI Endpoints)", level=1)

    H(doc, "18.1  What an endpoint is", level=2)
    P(doc,
      "An always-on HTTPS service that hosts one or more model versions. You POST a JSON "
      "instance, it returns a prediction. Vertex handles autoscaling, GPU attach, traffic "
      "split between versions, regional endpoints, IAM auth."
    )

    H(doc, "18.2  Machine types per detector", level=2)
    BULLET(doc, "rcf-metrics-{env}      → n1-standard-2.")
    BULLET(doc, "iforest-logs-{env}     → n1-standard-2.")
    BULLET(doc, "lstm-ae-traces-{env}   → n1-standard-4 + 1× L4 GPU.")
    BULLET(doc, "log-embedding-{env}    → n1-standard-4 + 1× L4 GPU.")

    H(doc, "18.3  Traffic split for safe rollout", level=2)
    P(doc,
      "When a new model version is deployed, you can split traffic 90/10 between old and new. "
      "Watch metrics for an hour; promote to 100/0 or roll back. This is canary deployment "
      "for ML."
    )

    MAP_BOX(doc,
            "Endpoints in this repo",
            "infra/modules/vertex/main.tf — 4 google_vertex_ai_endpoint resources",
            "Endpoints are pre-created (empty) by Terraform; Stage 6 of deploy_all.sh attaches "
            "the trained model to the matching endpoint. This split is intentional — endpoint "
            "lifecycle is infra, model lifecycle is ML.")

    STITCH(doc,
           "a stable URL per detector that returns predictions",
           "the orchestration service that calls each endpoint, picks the winner, stores the "
           "alert, and surfaces it for humans — the Scoring API")

    doc.add_page_break()

    # ── Chapter 19 — Scoring API ───────────────────────────────────────
    H(doc, "Chapter 19 — The Scoring API — wiring detectors into one HTTP surface", level=1)

    H(doc, "19.1  Why a thin gateway in front of the endpoints", level=2)
    BULLET(doc, "Endpoints are private; gateway adds public routing + Cloud Armor.")
    BULLET(doc, "Decide which detector to call based on event source.")
    BULLET(doc, "Apply per-detector thresholds.")
    BULLET(doc, "Persist alerts (Firestore) for the UI.")
    BULLET(doc, "Receive feedback labels and write back.")

    H(doc, "19.2  The five HTTP routes", level=2)
    BULLET(doc, "POST /api/v1/score             — single-event scoring.")
    BULLET(doc, "GET  /api/v1/alerts            — list alerts (filters: since/source/severity/limit).")
    BULLET(doc, "GET  /api/v1/alerts/{id}/explain — feature contributions for one alert.")
    BULLET(doc, "POST /api/v1/feedback           — true_positive / false_positive label.")
    BULLET(doc, "GET  /api/v1/sources            — last-seen heartbeat per source.")

    H(doc, "19.3  Internals", level=2)
    BULLET(doc, "FastAPI on uvicorn, Gunicorn-managed, autoscaled by HPA on CPU.")
    BULLET(doc, "Vertex Endpoint client wraps aiplatform.Endpoint(name).predict(instances=...).")
    BULLET(doc, "Firestore is the alert store — schemaless, horizontally scalable, free tier covers demo.")
    BULLET(doc, "Prometheus /metrics endpoint exposes scoring_requests_total + scoring_latency_seconds.")

    MAP_BOX(doc,
            "Scoring API",
            "api/scoring/main.py + services/{vertex_client,store,telemetry}.py",
            "Helm chart helm/charts/anomaly-scoring-api deploys it with Ingress + ManagedCertificate + Cloud Armor.")

    STITCH(doc,
           "an HTTP service that produces alerts",
           "the human surface those alerts surface in — the AIOps Console UI")

    doc.add_page_break()

    # ── Chapter 20 — UI ────────────────────────────────────────────────
    H(doc, "Chapter 20 — The AIOps Console — the human-facing UI", level=1)

    H(doc, "20.1  Why a custom UI", level=2)
    P(doc,
      "Cloud Monitoring dashboards are great for time-series. They are weak at: per-alert "
      "drill-in, feedback loops, detector explanations, custom severity views. So we ship a "
      "small static console that calls the Scoring API directly."
    )

    H(doc, "20.2  What it shows", level=2)
    BULLET(doc, "Live alert table with filters (source / severity / limit).")
    BULLET(doc, "Stats cards (last 1h: total / high / medium / detector count).")
    BULLET(doc, "Source-health sidebar (lag in seconds per source).")
    BULLET(doc, "Per-alert 'Explain' modal — top features, observed vs baseline.")
    BULLET(doc, "Feedback buttons (Mark real / Mark false positive) → POST /feedback.")

    H(doc, "20.3  Architecture", level=2)
    P(doc,
      "Plain HTML + JS. Hosted by NGINX in api/ui/. The same NGINX proxies /api/* to the "
      "ClusterIP Service of anomaly-scoring-api. Deployed by helm/charts/aiops-ui with a "
      "GCE Ingress + ManagedCertificate + Cloud Armor."
    )

    MAP_BOX(doc,
            "Console code",
            "api/ui/site/index.html  · api/ui/nginx.conf  · helm/charts/aiops-ui/",
            "Open https://$DOMAIN_AIOPS after deploy.")

    STITCH(doc,
           "a working alert console",
           "the lifecycle that keeps detectors useful as the world changes — drift, retraining, "
           "and the closed loop")

    doc.add_page_break()

    # ── Chapter 21 — Drift ─────────────────────────────────────────────
    H(doc, "Chapter 21 — Drift, retraining, and the closed loop", level=1)

    H(doc, "21.1  Drift", level=2)
    P(doc,
      "Models age. Two failure modes:"
    )
    BULLET(doc, "Data drift — input distribution changes (your traffic mix shifts; a new feature is rolled out).")
    BULLET(doc, "Concept drift — relationship between input and label changes (what was 'normal' yesterday isn't today).")

    H(doc, "21.2  Detect drift on inputs, not predictions (CLAUDE.md rule 5)", level=2)
    P(doc,
      "It's much faster to detect that req_count_5m's mean has shifted by 3σ than to wait for "
      "a precision dip. Compute statistical distance (KL divergence, KS test) between current "
      "feature distribution and a frozen reference distribution from the training run."
    )

    H(doc, "21.3  Retraining cadence", level=2)
    BULLET(doc, "RCF / IForest — daily (cheap, fast).")
    BULLET(doc, "LSTM-AE / LogBERT — weekly (expensive, GPU).")
    BULLET(doc, "Manual trigger — drift alarm or operator publish to retrain topic.")

    MAP_BOX(doc,
            "Schedules + alarm",
            "infra/modules/monitoring/main.tf — 4 google_cloud_scheduler_job + google_monitoring_alert_policy.precision_drop",
            "Schedulers publish to retrain Pub/Sub topic; a Cloud Function consumes and re-runs the matching pipeline.")

    H(doc, "21.4  Feedback loop", level=2)
    P(doc,
      "Every UI feedback click writes a row to Firestore (alert_id, label). Future training "
      "runs read those labels and use them as ground truth for evaluation, weighting, and (for "
      "supervised refinement) labelled training. This is what makes the platform learn."
    )

    STITCH(doc,
           "a closed lifecycle that keeps detectors current",
           "the security posture that protects every layer of the system — defence in depth")

    doc.add_page_break()

    # ── Chapter 22 — Security ──────────────────────────────────────────
    H(doc, "Chapter 22 — Security across the stack — defence in depth", level=1)

    P(doc, "Multiple weak walls > one strong wall. Six controls in this repo:")

    NUM(doc, "Cloud Armor at the edge — WAF rulesets (sqli, xss) + 600 req/min/IP rate-limit.")
    NUM(doc, "ManagedCertificate + HTTPS-only — TLS terminates at the global LB.")
    NUM(doc, "Workload Identity for every pod — no static service-account keys in cluster.")
    NUM(doc, "Secret Manager — passwords/keys versioned, IAM-controlled, never in YAML or git.")
    NUM(doc, "NetworkPolicy — pod-level firewall, blocks egress to GCE metadata server (169.254.169.254).")
    NUM(doc, "Audit logs — Cloud Audit Logs capture every IAM call automatically.")

    P(doc, "PII rule (CLAUDE.md rule 1):", bold=True)
    P(doc, "All usernames + IPs are HMAC-hashed inside the parser. Detectors only see hashes. "
            "You cannot leak data you do not store.")

    MAP_BOX(doc,
            "Cloud Armor policy",
            "infra/modules/lb/main.tf — google_compute_security_policy.armor",
            "Three rule families: rate_limit, sqli WAF, xss WAF. Edit carefully — public blast radius.")

    STITCH(doc,
           "a layered security posture",
           "the pragmatic side of running this on a credit card — what actually drives the bill, "
           "and which knobs lower it")

    doc.add_page_break()

    # ── Chapter 23 — Cost ─────────────────────────────────────────────
    H(doc, "Chapter 23 — Cost engineering — knobs that decide your bill", level=1)

    H(doc, "23.1  Where the money goes (dev environment)", level=2)
    t = doc.add_table(rows=1, cols=4); t.style = "Light Grid Accent 1"
    head = t.rows[0].cells
    for ci, h in enumerate(["Resource", "Type", "Hourly", "Used"]):
        head[ci].text = h; shade(head[ci], "DCE6F1")
        for p in head[ci].paragraphs:
            for r in p.runs: r.bold = True
    rows = [
        ("Train (RCF/IForest)",      "n1-standard-4 Spot",        "≈ $0.05",   "daily, 10-30 min"),
        ("Train (LSTM-AE/LogBERT)",  "g2-standard-8 + L4 Spot",   "≈ $0.30",   "weekly, 1-3 h"),
        ("Serve (RCF/IForest)",      "n1-standard-2",             "≈ $0.10",   "continuous"),
        ("Serve (LSTM-AE/LogBERT)",  "n1-standard-4 + L4",        "≈ $0.95",   "continuous"),
        ("GKE Autopilot",            "per-pod",                   "~$0.04/pod","2-8 replicas"),
        ("Cloud SQL",                "db-custom-1-3840",          "≈ $0.07",   "continuous"),
        ("GCS / Pub/Sub / Logging",  "volume-based",              "< $5/day",  "continuous"),
        ("BigQuery queries",         "$5 / TB scanned",           "varies",    "per query"),
    ]
    for r in rows:
        cells = t.add_row().cells
        for ci, val in enumerate(r):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for run in p.runs: run.font.size = Pt(10)
    doc.add_paragraph()

    H(doc, "23.2  Cost rules we follow (CLAUDE.md rule 2)", level=2)
    BULLET(doc, "Spot for all training (3-4× cheaper, preempts mid-job → restart logic in pipelines).")
    BULLET(doc, "L4 GPU is the ceiling. No A100 / H100 in this project.")
    BULLET(doc, "n1-standard-2 endpoints in dev. Only scale up serving for prod.")
    BULLET(doc, "Stop Cloud SQL during off-hours.")
    BULLET(doc, "GCS lifecycle rule deletes raw partitions after 90 days.")
    BULLET(doc, "BigQuery: WHERE date_partition = ... before SELECT — never scan the whole table.")

    STITCH(doc,
           "a mental model of where dollars go and how to control them",
           "the actual button-press to bring the whole stack up — the deploy walk-through")

    doc.add_page_break()

    # ── Chapter 24 — Deploy walk ────────────────────────────────────────
    H(doc, "Chapter 24 — The full deploy walk-through (scripts/deploy_all.sh)", level=1)

    P(doc,
      "Single command, 9 stages, interactive. Each stage prints a label and asks "
      "[Y]es / [n]o-skip / [s]hell / [q]uit before running. Outputs from earlier stages "
      "are cached in .deploy.env so re-runs reuse them."
    )

    stages = [
        ("1", "Platform Terraform (infra/)",
         "Creates: GKE Autopilot, GCS bucket, Pub/Sub topics, Artifact Registry, Cloud SQL Postgres, "
         "Cloud Armor policy, BigQuery dataset + 3 tables, 4 Vertex Endpoints (empty), "
         "Cloud Scheduler retrain jobs, 2 Cloud Monitoring dashboards. Takes ~10 min the first time "
         "(Autopilot cluster bring-up is the slow bit)."),
        ("2", "Build + push images",
         "Builds Docker images for: scoring-api, ui, 4 trainers, demo-app trio. Pushes to "
         "Artifact Registry. Re-tags Vertex pre-built sklearn image as serve image for the "
         "two CPU detectors."),
        ("3", "Demo-app Terraform",
         "Cloud SQL MySQL with slow-log → Pub/Sub sink. Memorystore Redis. demo-api/demo-worker GSAs "
         "with WLI bindings + Secret Accessor IAM. Static IP demo-app-ip."),
        ("4", "Streaming Cloud Function",
         "Deploys ml/streaming/ (entry-point handler) to Cloud Functions Gen2. Subscribes to "
         "events Pub/Sub topic. Emits anomalies to anomalies topic + Cloud Monitoring + Cloud Logging."),
        ("5", "Feature engineering + 4 Vertex Pipelines",
         "Seeds 5000 synthetic events to Pub/Sub. Builds 5-min features in GCS. Submits 4 KFP "
         "pipelines to Vertex. Each registers a model on success."),
        ("6", "Attach trained models to Vertex Endpoints",
         "Looks up newest registered model per display name. Deploys to matching endpoint with "
         "right machine type. Sets traffic split 100% → new version."),
        ("7", "Helm: scoring-api + aiops-ui + otel-collector",
         "Three Helm releases on GKE. Scoring API gets endpoint IDs as env vars. UI gets DOMAIN_AIOPS. "
         "OTEL collector runs in observability namespace, exports to Cloud Trace + GMP."),
        ("8", "Helm: demo-api + demo-worker + demo-web",
         "Pulls SQL password from Secret Manager into k8s Secret demo-api-db. Deploys demo trio with "
         "WLI annotations + DB host + Redis URL."),
        ("9", "Traffic-gen CronJob + smoke test",
         "Applies the Locust CronJob (image ref substituted with project). Kicks one job immediately. "
         "Runs scripts/smoke_test.sh against $DOMAIN_AIOPS."),
    ]
    t = doc.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
    head = t.rows[0].cells
    for ci, h in enumerate(["Stage", "Title", "What it does"]):
        head[ci].text = h; shade(head[ci], "DCE6F1")
        for p in head[ci].paragraphs:
            for r in p.runs: r.bold = True
    for s in stages:
        cells = t.add_row().cells
        for ci, val in enumerate(s):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(10)

    P(doc, "End-state surfaces", bold=True)
    BULLET(doc, "https://$DOMAIN_AIOPS — AIOps Anomaly Console.")
    BULLET(doc, "https://$DOMAIN_AIOPS/api/v1/alerts — Scoring API JSON.")
    BULLET(doc, "Cloud Monitoring dashboards — AIOps Overview + Detector Health.")
    BULLET(doc, "BigQuery dataset 'monitoring' — raw_events, features_security, anomalies.")
    BULLET(doc, "Pub/Sub anomalies-fanout subscription — pull external consumers from here.")
    BULLET(doc, "Cloud Trace — OTEL spans from every demo-app pod.")

    STITCH(doc,
           "a deployed stack",
           "the most likely things to go wrong while you operate it, and how to recover")

    doc.add_page_break()

    # ── Chapter 25 — Failure modes ─────────────────────────────────────
    H(doc, "Chapter 25 — Failure modes you will hit, and how to fix them", level=1)

    failures = [
        ("Cluster create fails — quota",
         "You haven't raised L4 / vCPU quota in the region.",
         "GCP Console → IAM & Admin → Quotas → request increase for region asia-south1."),
        ("Vertex Pipeline fails — 'no images'",
         "Trainer image not pushed.",
         "Re-run Stage 2; check Artifact Registry."),
        ("Endpoint deploy fails — 'no model'",
         "Pipeline didn't register a model (gate failed or earlier crash).",
         "Look in Vertex Console → Pipelines → run; fix gate or features and re-run."),
        ("Helm install pending — Ingress no IP",
         "ManagedCertificate provisioning takes 15-60 min on first request.",
         "Wait. Check `kubectl describe managedcertificate aiops-ui-cert`."),
        ("UI shows 502",
         "Scoring API not Ready.",
         "`kubectl get pods` — check logs. Most often the Vertex endpoint env var is empty."),
        ("Streaming Function not firing",
         "Pub/Sub subscription not created or wrong topic.",
         "`gcloud pubsub topics list` + `gcloud functions describe`. Verify trigger."),
        ("Cloud SQL private IP unreachable",
         "PSA peering missing or network mismatch.",
         "Apply parent infra first. Confirm google_service_networking_connection.psa exists."),
        ("OTEL traces missing in Cloud Trace",
         "Workload Identity annotation on otel-collector SA missing.",
         "Patch SA: kubectl annotate sa -n observability otel-collector iam.gke.io/gcp-service-account=$RUNNER_SA"),
        ("BigQuery scans cost a fortune",
         "WHERE missing on partition column.",
         "Always include `WHERE _PARTITIONTIME >= ...` on native, hive partition WHERE on external."),
        ("Worker pod CrashLoopBackOff",
         "Redis URL wrong or Memorystore not in same VPC.",
         "Check REDIS_URL env. Confirm authorized_network is 'default'."),
    ]
    t = doc.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
    head = t.rows[0].cells
    for ci, h in enumerate(["Symptom", "Root cause", "Fix"]):
        head[ci].text = h; shade(head[ci], "DCE6F1")
        for p in head[ci].paragraphs:
            for r in p.runs: r.bold = True
    for f in failures:
        cells = t.add_row().cells
        for ci, val in enumerate(f):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph()

    STITCH(doc,
           "operational confidence",
           "every term used in this guide, in one alphabetical place — the glossary")

    doc.add_page_break()

    # ── Chapter 26 — Glossary ──────────────────────────────────────────
    H(doc, "Chapter 26 — Glossary", level=1)
    glossary = [
        ("Anomaly", "A data point unlike the majority — usually a rare or fresh pattern."),
        ("Artifact Registry", "GCP-managed Docker registry."),
        ("Autopilot (GKE)", "GKE mode where Google manages the nodes; you only declare pods."),
        ("BackendConfig", "GKE CRD that attaches Cloud Armor or custom health-check to a Service."),
        ("BigQuery external table", "SQL table whose data lives in GCS, read on demand."),
        ("CommonEvent", "Project-internal canonical event schema (ml/parsers/__init__.py)."),
        ("Counter / Gauge / Histogram", "The three Prometheus metric shapes."),
        ("DAG", "Directed Acyclic Graph — pipelines are DAGs of steps."),
        ("Dataflow", "Managed Apache Beam runner on GCP."),
        ("Drift (data / concept)", "Inputs or input-output relationship changing over time."),
        ("EWMA", "Exponentially Weighted Moving Average."),
        ("Endpoint (Vertex)", "Always-on HTTPS URL hosting one or more model versions."),
        ("Feature engineering", "Turning raw events into model-ready columns."),
        ("Flex template (Dataflow)", "Containerised Beam pipeline you launch with parameters."),
        ("GMP", "Google Managed Prometheus."),
        ("Helm chart", "Templated Kubernetes manifest bundle."),
        ("Hive partitioning", "Directory layout source=app/ that tells query engines to skip irrelevant files."),
        ("HMAC", "Keyed hash; what we use to scrub PII without losing distinct-count utility."),
        ("HPA", "HorizontalPodAutoscaler; scales pods on CPU/custom metrics."),
        ("IAM", "Identity & Access Management. Roles, bindings, principals, scopes."),
        ("Ingress", "Kubernetes object that maps external HTTP routes to Services."),
        ("KFP v2", "Kubeflow Pipelines v2 — Vertex Pipelines DSL."),
        ("Lake / Warehouse", "Lake = raw files. Warehouse = SQL-optimised store."),
        ("Log sink", "Cloud Logging routing rule — exports filtered logs to a destination."),
        ("ManagedCertificate", "GKE CRD for Google-issued TLS certs."),
        ("MLOps", "DevOps for ML — versioning, deployment, monitoring of models."),
        ("Multi-stage build", "Dockerfile pattern — compile in one stage, copy artefact to a smaller stage."),
        ("OTEL / OTLP", "OpenTelemetry SDK / wire format for telemetry."),
        ("PodMonitoring", "GMP CRD declaring which pods to scrape for /metrics."),
        ("Precision / Recall / F1 / AUC / P@1%", "Classification quality metrics."),
        ("PSA", "Private Service Access — VPC peering for Google-managed services."),
        ("Pub/Sub", "GCP pub-sub message bus."),
        ("Quota", "GCP-imposed cap on resource counts; raise via Console."),
        ("Random Cut Forest", "Ensemble of random trees used for streaming anomaly detection."),
        ("Rolling window / Tumbling window", "Sliding (overlapping) vs fixed-non-overlapping aggregation windows."),
        ("Scoring API", "FastAPI gateway over Vertex Endpoints."),
        ("Score", "Numeric output of a detector; thresholded into 'anomaly'."),
        ("Service / Endpoint / Ingress", "K8s Service = internal DNS. Ingress = external. Endpoint = Vertex serving."),
        ("Spot VM", "Preemptible VM at ~3× discount."),
        ("StatefulSet", "K8s controller for stateful workloads (not used here; Deployments only)."),
        ("Threshold-based / ML-based", "Detector tier 2 vs tier 3."),
        ("Trace / Span / Trace ID", "OTEL data model — tree of spans tied by ID."),
        ("Vertex AI", "GCP umbrella for ML training, registry, serving, pipelines."),
        ("VPC", "Virtual Private Cloud — private network in GCP."),
        ("WAF", "Web Application Firewall (Cloud Armor)."),
        ("Workload Identity", "Bind a KSA to a GSA without static keys."),
        ("Z-score", "(x - mean) / stddev. Bigger absolute value = rarer point."),
    ]
    t = doc.add_table(rows=1, cols=2); t.style = "Light Grid Accent 1"
    head = t.rows[0].cells
    head[0].text, head[1].text = "Term", "Definition"
    for c in head:
        shade(c, "DCE6F1")
        for p in c.paragraphs:
            for r in p.runs: r.bold = True
    for term, defn in glossary:
        cells = t.add_row().cells
        cells[0].text, cells[1].text = term, defn
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(10)

    SECTION_BREAK(doc)

    # ── Chapter 27 — Study plan ─────────────────────────────────────────
    H(doc, "Chapter 27 — A 4-week study plan", level=1)

    H(doc, "Week 1 — Foundations", level=2)
    BULLET(doc, "Day 1-2: Read this guide chapters 1-5. Open every file mentioned in MAP_BOX entries.")
    BULLET(doc, "Day 3: Install gcloud + terraform + kubectl + helm + docker locally. Auth: gcloud auth login.")
    BULLET(doc, "Day 4: Build the demo-api image locally; run with docker compose against a local MySQL.")
    BULLET(doc, "Day 5: Run scripts/deploy_all.sh stages 1-3 in a sandbox project. Watch resources appear in the GCP Console.")

    H(doc, "Week 2 — Data + ML", level=2)
    BULLET(doc, "Day 1: Read chapters 6-11. Trace one synthetic event from demo-api → Cloud Logging → Pub/Sub → GCS.")
    BULLET(doc, "Day 2: Open BigQuery; query dataset monitoring.raw_events. Filter by source.")
    BULLET(doc, "Day 3: Read chapters 12-14. Run Stage 5 — submit pipelines. Watch them in Vertex Console.")
    BULLET(doc, "Day 4: Open ml/pipelines/rcf_metrics/pipeline.py. Read every component step.")
    BULLET(doc, "Day 5: Modify the pipeline — add a histogram of feature values to the eval step. Re-run.")

    H(doc, "Week 3 — Serving + UI + AIOps", level=2)
    BULLET(doc, "Day 1: Read chapters 15-20. Run Stages 6-9.")
    BULLET(doc, "Day 2: Hit /api/v1/alerts with curl. POST a synthetic event to /api/v1/score.")
    BULLET(doc, "Day 3: Open the AIOps Console. Trigger an attack — `kubectl create job --from=cronjob/demo-traffic-attack`.")
    BULLET(doc, "Day 4: Click Explain on a real alert. Read api/scoring/main.py — follow the call path.")
    BULLET(doc, "Day 5: Add a new severity badge or a new column to the UI table.")

    H(doc, "Week 4 — Hardening + extending", level=2)
    BULLET(doc, "Day 1: Read chapters 21-23. Trigger a manual retrain by publishing to retrain Pub/Sub.")
    BULLET(doc, "Day 2: Add a 5th detector — simple seasonal-decompose on metrics. New folder, new endpoint.")
    BULLET(doc, "Day 3: Add a Slack notifier — Cloud Function on the anomalies topic.")
    BULLET(doc, "Day 4: Tighten Cloud Armor — restrict by source IP. Test by running locust from a blocked IP.")
    BULLET(doc, "Day 5: Run scripts/teardown.sh. Recreate from scratch. The second deploy should be muscle memory.")

    H(doc, "Recommended reading", level=2)
    BULLET(doc, "Designing Machine Learning Systems — Chip Huyen.")
    BULLET(doc, "The Site Reliability Workbook — Google. Chapter on observability is gold.")
    BULLET(doc, "Coursera — MLOps Specialization (Andrew Ng + Robert Crowe).")
    BULLET(doc, "Google Cloud Skill Boosts — 'Build and Deploy ML Solutions on Vertex AI'.")
    BULLET(doc, "Kubernetes Up & Running — Kelsey Hightower.")
    BULLET(doc, "OpenTelemetry's docs site, otel.io — read the data-model page.")

    P(doc,
      "When you finish week 4 you will not just have read about MLOps — you will have built, "
      "broken, retrained, and shipped one. That is the only way to learn this discipline.",
      italic=True)

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
